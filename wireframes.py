"""
Wireframe Generator - Klantonderhoudsysteem
Reverse engineered wireframes uit de PHP bestanden
Genereert een Word document met wireframes van elke pagina
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# Pagina marges kleiner maken
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)


def stel_cel_kleur_in(cel, kleur):
    """Achtergrondkleur van een tabelcel instellen"""
    shading = cel._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): kleur,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elem)


def maak_rand(tabel):
    """Rand om een tabel zetten"""
    tbl = tabel._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for rand in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        element = borders.makeelement(qn(f'w:{rand}'), {
            qn('w:val'): 'single',
            qn('w:sz'): '4',
            qn('w:space'): '0',
            qn('w:color'): '999999'
        })
        borders.append(element)
    tblPr.append(borders)


def maak_dikke_rand(tabel):
    """Dikke rand om een tabel (voor wireframe container)"""
    tbl = tabel._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for rand in ['top', 'left', 'bottom', 'right']:
        element = borders.makeelement(qn(f'w:{rand}'), {
            qn('w:val'): 'single',
            qn('w:sz'): '12',
            qn('w:space'): '0',
            qn('w:color'): '333333'
        })
        borders.append(element)
    for rand in ['insideH', 'insideV']:
        element = borders.makeelement(qn(f'w:{rand}'), {
            qn('w:val'): 'none',
            qn('w:sz'): '0',
            qn('w:space'): '0',
            qn('w:color'): 'FFFFFF'
        })
        borders.append(element)
    tblPr.append(borders)


def voeg_lege_regel_toe():
    doc.add_paragraph("")


def voeg_invoerveld_toe(cel, label_tekst, breedte="breed", waarde=""):
    """Voegt een label + invoerveld toe in een cel"""
    p = cel.add_paragraph()
    run = p.add_run(label_tekst)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    p2 = cel.add_paragraph()
    veld_tekst = waarde if waarde else "________________________"
    if breedte == "breed":
        veld_tekst = waarde if waarde else "________________________________________"
    run2 = p2.add_run(f"[  {veld_tekst}  ]")
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


# ============================================================
# TITELPAGINA
# ============================================================
doc.add_paragraph("")
doc.add_paragraph("")
doc.add_paragraph("")

titel = doc.add_paragraph()
titel.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = titel.add_run("Wireframes")
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

subtitel = doc.add_paragraph()
subtitel.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitel.add_run("Klantonderhoudsysteem")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x4C, 0xAF, 0x50)

doc.add_paragraph("")

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("Reverse engineered uit het PHP project")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph("")
doc.add_paragraph("")

# Pagina overzicht
overzicht = doc.add_paragraph()
overzicht.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = overzicht.add_run("Pagina overzicht:")
run.font.size = Pt(12)
run.bold = True

paginas = [
    "1. Login pagina (index.php)",
    "2. Admin Dashboard (admin.php)",
    "3. Klant pagina (klant.php)",
    "4. Nieuwe klant toevoegen (toevoegen.php)",
    "5. Klant bewerken (bewerken.php)",
    "6. Verwijder bevestiging (verwijderen.php)",
    "7. Navigatie flow diagram"
]

for pagina in paginas:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(pagina)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()


# ============================================================
# WIREFRAME 1: LOGIN PAGINA (index.php)
# ============================================================
titel = doc.add_heading("Wireframe 1: Login Pagina", level=1)
titel.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

p = doc.add_paragraph()
run = p.add_run("Bestand: index.php")
run.font.size = Pt(9)
run.italic = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

p = doc.add_paragraph()
run = p.add_run("Beschrijving: ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("De landingspagina van het systeem. Gebruikers loggen hier in als admin of klant.")
run.font.size = Pt(10)

voeg_lege_regel_toe()

# Wireframe container
container = doc.add_table(rows=1, cols=1)
container.alignment = WD_TABLE_ALIGNMENT.CENTER
maak_dikke_rand(container)
cel = container.cell(0, 0)
stel_cel_kleur_in(cel, 'F5F5F5')

# Browser bar
p = cel.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("  🌐  localhost/lastig/index.php")
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# Titel
p = cel.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Klantonderhoudsysteem")
run.font.size = Pt(18)
run.bold = True

# Welkomsttekst
p = cel.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Welkom bij het klantonderhoudsysteem. Log in om verder te gaan.")
run.font.size = Pt(10)

cel.add_paragraph("")

# Foutmelding (optioneel)
p = cel.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("[Foutmelding: Verkeerde gebruikersnaam of wachtwoord!]")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
run.italic = True

cel.add_paragraph("")

# Login formulier
voeg_invoerveld_toe(cel, "        Gebruikersnaam:")
cel.add_paragraph("")
voeg_invoerveld_toe(cel, "        Wachtwoord:")
cel.add_paragraph("")

# Knop
p = cel.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("        [ Inloggen ]")
run.font.size = Pt(10)
run.bold = True
run.font.color.rgb = RGBColor(0x4C, 0xAF, 0x50)

cel.add_paragraph("")
cel.add_paragraph("")

# Toelichting
voeg_lege_regel_toe()
p = doc.add_paragraph()
run = p.add_run("Werking: ")
run.bold = True
run.font.size = Pt(9)
run = p.add_run("Formulier stuurt POST request naar login.php. Bij juiste gegevens: admin -> admin.php, klant -> klant.php. Bij fout: terug naar index.php met foutmelding.")
run.font.size = Pt(9)

doc.add_page_break()


# ============================================================
# WIREFRAME 2: ADMIN DASHBOARD (admin.php)
# ============================================================
titel = doc.add_heading("Wireframe 2: Admin Dashboard", level=1)
titel.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

p = doc.add_paragraph()
run = p.add_run("Bestand: admin.php")
run.font.size = Pt(9)
run.italic = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

p = doc.add_paragraph()
run = p.add_run("Beschrijving: ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("Het admin dashboard toont een overzicht van alle klanten in een tabel. De admin kan klanten toevoegen, bewerken en verwijderen.")
run.font.size = Pt(10)

voeg_lege_regel_toe()

# Container
container = doc.add_table(rows=1, cols=1)
container.alignment = WD_TABLE_ALIGNMENT.CENTER
maak_dikke_rand(container)
cel = container.cell(0, 0)
stel_cel_kleur_in(cel, 'F5F5F5')

# Browser bar
p = cel.paragraphs[0]
run = p.add_run("  🌐  localhost/lastig/admin.php")
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# Titel
p = cel.add_paragraph()
run = p.add_run("Admin Dashboard")
run.font.size = Pt(18)
run.bold = True

# Welkom + uitloggen
p = cel.add_paragraph()
run = p.add_run("Welkom, admin!  ")
run.font.size = Pt(10)
run = p.add_run("[Uitloggen]")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x4C, 0xAF, 0x50)
run.underline = True

cel.add_paragraph("")

# Subtitel
p = cel.add_paragraph()
run = p.add_run("Alle Klanten")
run.font.size = Pt(14)
run.bold = True

# Link toevoegen
p = cel.add_paragraph()
run = p.add_run("+ Nieuwe klant toevoegen")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x4C, 0xAF, 0x50)
run.underline = True

cel.add_paragraph("")

# Klanten tabel
tabel = doc.add_table(rows=4, cols=8)
tabel.alignment = WD_TABLE_ALIGNMENT.CENTER
maak_rand(tabel)

headers = ["ID", "Voornaam", "Achternaam", "Email", "Telefoon", "Adres", "Woonplaats", "Acties"]
for i, header in enumerate(headers):
    cel_h = tabel.cell(0, i)
    stel_cel_kleur_in(cel_h, '4CAF50')
    p = cel_h.paragraphs[0]
    run = p.add_run(header)
    run.font.size = Pt(8)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# Rij 1
data1 = ["1", "Jan", "de Vries", "jan@email.nl", "06-12345678", "Hoofdstraat 1", "Amsterdam", ""]
for i, d in enumerate(data1):
    p = tabel.cell(1, i).paragraphs[0]
    run = p.add_run(d)
    run.font.size = Pt(7)

# Acties rij 1
p = tabel.cell(1, 7).paragraphs[0]
run = p.add_run("Bewerken")
run.font.size = Pt(7)
run.font.color.rgb = RGBColor(0x4C, 0xAF, 0x50)
run = p.add_run(" ")
run = p.add_run("Verwijderen")
run.font.size = Pt(7)
run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

# Rij 2
data2 = ["2", "Piet", "Bakker", "piet@email.nl", "06-87654321", "Kerkweg 5", "Rotterdam", ""]
for i, d in enumerate(data2):
    p = tabel.cell(2, i).paragraphs[0]
    run = p.add_run(d)
    run.font.size = Pt(7)
stel_cel_kleur_in(tabel.cell(2, 0), 'F2F2F2')
stel_cel_kleur_in(tabel.cell(2, 1), 'F2F2F2')
stel_cel_kleur_in(tabel.cell(2, 2), 'F2F2F2')
stel_cel_kleur_in(tabel.cell(2, 3), 'F2F2F2')
stel_cel_kleur_in(tabel.cell(2, 4), 'F2F2F2')
stel_cel_kleur_in(tabel.cell(2, 5), 'F2F2F2')
stel_cel_kleur_in(tabel.cell(2, 6), 'F2F2F2')
stel_cel_kleur_in(tabel.cell(2, 7), 'F2F2F2')

# Acties rij 2
p = tabel.cell(2, 7).paragraphs[0]
run = p.add_run("Bewerken")
run.font.size = Pt(7)
run.font.color.rgb = RGBColor(0x4C, 0xAF, 0x50)
run = p.add_run(" ")
run = p.add_run("Verwijderen")
run.font.size = Pt(7)
run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

# Rij 3 (leeg voorbeeld)
data3 = ["...", "...", "...", "...", "...", "...", "...", "..."]
for i, d in enumerate(data3):
    p = tabel.cell(3, i).paragraphs[0]
    run = p.add_run(d)
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

voeg_lege_regel_toe()

# Toelichting
p = doc.add_paragraph()
run = p.add_run("Werking: ")
run.bold = True
run.font.size = Pt(9)
run = p.add_run("Sessie check: alleen rol=admin heeft toegang. Tabel toont alle records uit Klant_File. Even rijen hebben grijze achtergrond (zebra-striping via CSS). Bewerken linkt naar bewerken.php?id=X, Verwijderen naar verwijderen.php?id=X met JavaScript confirm().")
run.font.size = Pt(9)

doc.add_page_break()


# ============================================================
# WIREFRAME 3: KLANT PAGINA (klant.php)
# ============================================================
titel = doc.add_heading("Wireframe 3: Mijn Gegevens (Klant)", level=1)
titel.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

p = doc.add_paragraph()
run = p.add_run("Bestand: klant.php")
run.font.size = Pt(9)
run.italic = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

p = doc.add_paragraph()
run = p.add_run("Beschrijving: ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("De klant ziet alleen zijn/haar eigen gegevens en kan deze bewerken en opslaan.")
run.font.size = Pt(10)

voeg_lege_regel_toe()

# Container
container = doc.add_table(rows=1, cols=1)
container.alignment = WD_TABLE_ALIGNMENT.CENTER
maak_dikke_rand(container)
cel = container.cell(0, 0)
stel_cel_kleur_in(cel, 'F5F5F5')

# Browser bar
p = cel.paragraphs[0]
run = p.add_run("  🌐  localhost/lastig/klant.php")
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# Titel
p = cel.add_paragraph()
run = p.add_run("Mijn Gegevens")
run.font.size = Pt(18)
run.bold = True

# Welkom + uitloggen
p = cel.add_paragraph()
run = p.add_run("Welkom, klant1!  ")
run.font.size = Pt(10)
run = p.add_run("[Uitloggen]")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x4C, 0xAF, 0x50)
run.underline = True

cel.add_paragraph("")

# Succesmelding
p = cel.add_paragraph()
run = p.add_run("[Melding: Gegevens opgeslagen!]")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
run.italic = True

cel.add_paragraph("")

# Formulier
voeg_invoerveld_toe(cel, "        Voornaam:", waarde="Jan")
cel.add_paragraph("")
voeg_invoerveld_toe(cel, "        Achternaam:", waarde="de Vries")
cel.add_paragraph("")
voeg_invoerveld_toe(cel, "        Email:", waarde="jan@email.nl")
cel.add_paragraph("")
voeg_invoerveld_toe(cel, "        Telefoon:", waarde="06-12345678")
cel.add_paragraph("")
voeg_invoerveld_toe(cel, "        Adres:", waarde="Hoofdstraat 1")
cel.add_paragraph("")
voeg_invoerveld_toe(cel, "        Woonplaats:", waarde="Amsterdam")
cel.add_paragraph("")

# Knop
p = cel.add_paragraph()
run = p.add_run("        [ Opslaan ]")
run.font.size = Pt(10)
run.bold = True
run.font.color.rgb = RGBColor(0x4C, 0xAF, 0x50)

cel.add_paragraph("")

# Toelichting
voeg_lege_regel_toe()
p = doc.add_paragraph()
run = p.add_run("Werking: ")
run.bold = True
run.font.size = Pt(9)
run = p.add_run("Sessie check: alleen rol=klant heeft toegang. Haalt klantgegevens op via login_id. POST request update de gegevens in Klant_File. Na opslaan verschijnt groene bevestiging.")
run.font.size = Pt(9)

doc.add_page_break()


# ============================================================
# WIREFRAME 4: KLANT TOEVOEGEN (toevoegen.php)
# ============================================================
titel = doc.add_heading("Wireframe 4: Nieuwe Klant Toevoegen", level=1)
titel.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

p = doc.add_paragraph()
run = p.add_run("Bestand: toevoegen.php")
run.font.size = Pt(9)
run.italic = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

p = doc.add_paragraph()
run = p.add_run("Beschrijving: ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("De admin kan via dit formulier een nieuwe klant toevoegen aan het systeem.")
run.font.size = Pt(10)

voeg_lege_regel_toe()

# Container
container = doc.add_table(rows=1, cols=1)
container.alignment = WD_TABLE_ALIGNMENT.CENTER
maak_dikke_rand(container)
cel = container.cell(0, 0)
stel_cel_kleur_in(cel, 'F5F5F5')

# Browser bar
p = cel.paragraphs[0]
run = p.add_run("  🌐  localhost/lastig/toevoegen.php")
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# Titel
p = cel.add_paragraph()
run = p.add_run("Nieuwe Klant Toevoegen")
run.font.size = Pt(18)
run.bold = True

# Terug link
p = cel.add_paragraph()
run = p.add_run("< Terug naar overzicht")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x4C, 0xAF, 0x50)
run.underline = True

cel.add_paragraph("")

# Formulier
voeg_invoerveld_toe(cel, "        Voornaam: *")
cel.add_paragraph("")
voeg_invoerveld_toe(cel, "        Achternaam: *")
cel.add_paragraph("")
voeg_invoerveld_toe(cel, "        Email: *")
cel.add_paragraph("")
voeg_invoerveld_toe(cel, "        Telefoon:")
cel.add_paragraph("")
voeg_invoerveld_toe(cel, "        Adres:")
cel.add_paragraph("")
voeg_invoerveld_toe(cel, "        Woonplaats:")
cel.add_paragraph("")

# Knop
p = cel.add_paragraph()
run = p.add_run("        [ Toevoegen ]")
run.font.size = Pt(10)
run.bold = True
run.font.color.rgb = RGBColor(0x4C, 0xAF, 0x50)

cel.add_paragraph("")

# Toelichting
voeg_lege_regel_toe()
p = doc.add_paragraph()
run = p.add_run("Werking: ")
run.bold = True
run.font.size = Pt(9)
run = p.add_run("Alleen admin. POST request voert INSERT INTO Klant_File uit. Velden met * zijn verplicht (HTML required). Na toevoegen redirect naar admin.php.")
run.font.size = Pt(9)

doc.add_page_break()


# ============================================================
# WIREFRAME 5: KLANT BEWERKEN (bewerken.php)
# ============================================================
titel = doc.add_heading("Wireframe 5: Klant Bewerken", level=1)
titel.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

p = doc.add_paragraph()
run = p.add_run("Bestand: bewerken.php?id=1")
run.font.size = Pt(9)
run.italic = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

p = doc.add_paragraph()
run = p.add_run("Beschrijving: ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("De admin kan de gegevens van een bestaande klant bewerken. De huidige gegevens staan al ingevuld in het formulier.")
run.font.size = Pt(10)

voeg_lege_regel_toe()

# Container
container = doc.add_table(rows=1, cols=1)
container.alignment = WD_TABLE_ALIGNMENT.CENTER
maak_dikke_rand(container)
cel = container.cell(0, 0)
stel_cel_kleur_in(cel, 'F5F5F5')

# Browser bar
p = cel.paragraphs[0]
run = p.add_run("  🌐  localhost/lastig/bewerken.php?id=1")
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# Titel
p = cel.add_paragraph()
run = p.add_run("Klant Bewerken")
run.font.size = Pt(18)
run.bold = True

# Terug link
p = cel.add_paragraph()
run = p.add_run("< Terug naar overzicht")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x4C, 0xAF, 0x50)
run.underline = True

cel.add_paragraph("")

# Formulier met ingevulde waarden
voeg_invoerveld_toe(cel, "        Voornaam:", waarde="Jan")
cel.add_paragraph("")
voeg_invoerveld_toe(cel, "        Achternaam:", waarde="de Vries")
cel.add_paragraph("")
voeg_invoerveld_toe(cel, "        Email:", waarde="jan@email.nl")
cel.add_paragraph("")
voeg_invoerveld_toe(cel, "        Telefoon:", waarde="06-12345678")
cel.add_paragraph("")
voeg_invoerveld_toe(cel, "        Adres:", waarde="Hoofdstraat 1")
cel.add_paragraph("")
voeg_invoerveld_toe(cel, "        Woonplaats:", waarde="Amsterdam")
cel.add_paragraph("")

# Knop
p = cel.add_paragraph()
run = p.add_run("        [ Opslaan ]")
run.font.size = Pt(10)
run.bold = True
run.font.color.rgb = RGBColor(0x4C, 0xAF, 0x50)

cel.add_paragraph("")

# Toelichting
voeg_lege_regel_toe()
p = doc.add_paragraph()
run = p.add_run("Werking: ")
run.bold = True
run.font.size = Pt(9)
run = p.add_run("Alleen admin. Klant-ID via GET parameter. Huidige gegevens worden opgehaald en in de value-attributen gezet. POST request voert UPDATE uit op Klant_File WHERE id=X. Redirect naar admin.php.")
run.font.size = Pt(9)

doc.add_page_break()


# ============================================================
# WIREFRAME 6: VERWIJDEREN BEVESTIGING (verwijderen.php)
# ============================================================
titel = doc.add_heading("Wireframe 6: Klant Verwijderen", level=1)
titel.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

p = doc.add_paragraph()
run = p.add_run("Bestand: verwijderen.php?id=1")
run.font.size = Pt(9)
run.italic = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

p = doc.add_paragraph()
run = p.add_run("Beschrijving: ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("Verwijderen werkt via een JavaScript confirm() popup vanuit admin.php. Na bevestiging wordt de klant verwijderd.")
run.font.size = Pt(10)

voeg_lege_regel_toe()

# Container - admin pagina met popup
container = doc.add_table(rows=1, cols=1)
container.alignment = WD_TABLE_ALIGNMENT.CENTER
maak_dikke_rand(container)
cel = container.cell(0, 0)
stel_cel_kleur_in(cel, 'F5F5F5')

# Browser bar
p = cel.paragraphs[0]
run = p.add_run("  🌐  localhost/lastig/admin.php")
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

cel.add_paragraph("")

# Gesimuleerde popup
p = cel.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("╔══════════════════════════════════════╗")
run.font.size = Pt(10)
run.font.name = "Consolas"

p = cel.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("║                                                              ║")
run.font.size = Pt(10)
run.font.name = "Consolas"

p = cel.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("║     Weet je het zeker?                          ║")
run.font.size = Pt(10)
run.font.name = "Consolas"

p = cel.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("║                                                              ║")
run.font.size = Pt(10)
run.font.name = "Consolas"

p = cel.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("║       [ OK ]          [ Annuleren ]          ║")
run.font.size = Pt(10)
run.font.name = "Consolas"

p = cel.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("║                                                              ║")
run.font.size = Pt(10)
run.font.name = "Consolas"

p = cel.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("╚══════════════════════════════════════╝")
run.font.size = Pt(10)
run.font.name = "Consolas"

cel.add_paragraph("")
cel.add_paragraph("")

# Toelichting
voeg_lege_regel_toe()
p = doc.add_paragraph()
run = p.add_run("Werking: ")
run.bold = True
run.font.size = Pt(9)
run = p.add_run("JavaScript confirm() popup verschijnt bij klik op 'Verwijderen' in admin.php. Bij OK: verwijderen.php?id=X voert DELETE FROM Klant_File WHERE id=X uit. Bij Annuleren: niets gebeurt. Redirect terug naar admin.php.")
run.font.size = Pt(9)

doc.add_page_break()


# ============================================================
# WIREFRAME 7: NAVIGATIE FLOW
# ============================================================
titel = doc.add_heading("Navigatie Flow Diagram", level=1)
titel.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)

p = doc.add_paragraph()
run = p.add_run("Beschrijving: ")
run.bold = True
run.font.size = Pt(10)
run = p.add_run("Overzicht van hoe de pagina's met elkaar verbonden zijn.")
run.font.size = Pt(10)

voeg_lege_regel_toe()

# Flow diagram als tekst
flow_lines = [
    "                    ┌─────────────────┐",
    "                    │   index.php     │",
    "                    │  (Login pagina) │",
    "                    └────────┬────────┘",
    "                             │",
    "                       login.php",
    "                      (verwerking)",
    "                             │",
    "              ┌──────────────┼──────────────┐",
    "              │              │              │",
    "              ▼              │              ▼",
    "     ┌────────────────┐     │    ┌────────────────┐",
    "     │   admin.php    │     │    │   klant.php    │",
    "     │  (Dashboard)   │     │    │ (Mijn Gegevens)│",
    "     └───────┬────────┘     │    └────────────────┘",
    "             │              │",
    "    ┌────────┼────────┐     │",
    "    │        │        │     │",
    "    ▼        ▼        ▼     │",
    "┌────────┐┌────────┐┌──────┴───┐",
    "│toevoeg.││bewerkn.││verwijder.│",
    "│  php   ││  php   ││   php    │",
    "└────────┘└────────┘└──────────┘",
    "",
    "   Alle pagina's -> logout.php -> index.php",
]

for line in flow_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(line)
    run.font.size = Pt(9)
    run.font.name = "Consolas"

voeg_lege_regel_toe()

# Legenda
p = doc.add_paragraph()
run = p.add_run("Legenda:")
run.bold = True
run.font.size = Pt(10)

legenda = [
    ("index.php", "Landingspagina met login formulier"),
    ("login.php", "Verwerkt login, stuurt door op basis van rol"),
    ("admin.php", "Admin dashboard met klantenoverzicht tabel"),
    ("klant.php", "Klant pagina - alleen eigen gegevens zien/bewerken"),
    ("toevoegen.php", "Formulier om nieuwe klant toe te voegen (admin)"),
    ("bewerken.php", "Formulier om klant te bewerken (admin)"),
    ("verwijderen.php", "Verwijdert klant uit database (admin)"),
    ("logout.php", "Sessie vernietigen, terug naar login"),
    ("db.php", "Database connectie (wordt ge-include)"),
    ("style.css", "CSS styling voor alle pagina's"),
]

for bestand, beschrijving in legenda:
    p = doc.add_paragraph()
    run = p.add_run(f"  {bestand}")
    run.font.size = Pt(9)
    run.bold = True
    run.font.name = "Consolas"
    run = p.add_run(f"  -  {beschrijving}")
    run.font.size = Pt(9)


# ============================================================
# OPSLAAN
# ============================================================
uitvoer_pad = os.path.join(os.path.dirname(__file__), "Wireframes_Klantonderhoudsysteem.docx")
doc.save(uitvoer_pad)
print(f"Wireframes opgeslagen als: {uitvoer_pad}")
print("Klaar!")
